#!/usr/bin/env python3
from __future__ import annotations

import argparse
import json
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


def add_table(doc: Document, headers: list[str], rows: list[list[object]]):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = str(header)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for run in cell.paragraphs[0].runs:
            run.bold = True
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = str(value)
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    return table


def fmt_ms(value: float | int | None) -> str:
    if not value:
        return "-"
    return f"{float(value):.0f}"


def fmt_1(value: float | int | None) -> str:
    if not value:
        return "-"
    return f"{float(value):.1f}"


def avg(values: list[float]) -> float:
    values = [v for v in values if v]
    return sum(values) / len(values) if values else 0.0


def scenario_kind(name: str) -> str:
    return "短文本" if "短文本" in name else "长文本"


def build_doc(data: dict) -> Document:
    doc = Document()
    doc.styles["Normal"].font.name = "Arial"
    doc.styles["Normal"].font.size = Pt(10.5)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(data.get("poc", {}).get("report_title") or "模型压测报告")
    run.bold = True
    run.font.size = Pt(20)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("TraceNex / Fy-api Channel Benchmark").font.size = Pt(11)

    model_count = len(data.get("models", []))
    intro = doc.add_paragraph()
    intro.add_run("TraceNex 是企业级 AI 多模型服务平台，集成多家主流大模型，提供统一 API 调用入口，支持按价格、性能、通道、场景等维度进行路由，帮助企业优化模型匹配，获得稳定、安全且高性价比的大模型使用体验。")
    doc.add_paragraph(
        f"本报告针对 {model_count} 款国产 SOTA 模型进行短文本并发与长文本基线性能测试。"
    )

    models = [m.get("model", "") for m in data.get("models", [])]
    info_rows = [
        ["平台名称", data.get("poc", {}).get("platform_name", "TraceNex")],
        ["模型名称与版本", ", ".join(models)],
        ["测评类型", "LLM 简单压测"],
        ["测评范围", data.get("poc", {}).get("test_scope", "")],
        ["报告编号", data.get("poc", {}).get("report_id", "")],
        ["测评时间", data.get("generated_at", "")],
        ["测试网关", data.get("gateway", "")],
    ]

    doc.add_heading("一、测评基本信息", level=1)
    add_table(doc, ["字段名称", "内容"], info_rows)

    doc.add_heading("二、测评环境", level=1)
    add_table(doc, ["环境类别", "参数配置", "备注"], [
        ["硬件配置", "执行端本地开发机；网关为 CN 生产环境", "实际模型推理由供应商渠道完成"],
        ["CPU核心数、内存容量", "未固定", "本次压测瓶颈主要在远端网关/供应商响应"],
        ["存储类型与容量", "本地文件输出", "生成 JSON / CSV / Markdown / DOCX 报告"],
        ["测评工具", "fy-poc-loadtest", "TraceNex channel-benchmark 工具"],
        ["语言/框架", "Python / httpx / python-docx", "OpenAI-compatible streaming 请求"],
    ])

    scenario_rows: list[list[object]] = []
    for model in data.get("models", [])[:1]:
        for channel in model.get("channels", [])[:1]:
            for scenario in channel.get("scenarios", []):
                meta = scenario.get("scenario", {})
                levels = scenario.get("levels", [])
                scenario_rows.append([
                    ", ".join(models),
                    meta.get("name", ""),
                    f"{meta.get('input_tokens', '-') } tokens",
                    ", ".join(str(level.get("concurrency")) for level in levels),
                    "长文本仅做 C=1 基线" if "长文本" in meta.get("name", "") else "完整覆盖 1/20/80/200 并发",
                ])

    doc.add_heading("三、测评内容结果", level=1)
    doc.add_paragraph("测评基本信息：")
    add_table(doc, ["测评对象", "测评场景", "输入规格", "测试并发", "说明"], scenario_rows)

    doc.add_heading("四、测试方法", level=1)
    for text in [
        "本次测试使用 fy-poc-loadtest 工具，通过 OpenAI-compatible /v1/chat/completions 接口发起真实流式请求。",
        "短文本场景用于并发能力验证，覆盖并发 1、20、80、200。",
        "长文本 7K 场景仅做 C=1 基线测试，不覆盖长文本高并发。",
        "关键指标包括 TTFT、Latency、TPOT、tokens/s、请求成功率和错误分布。",
    ]:
        doc.add_paragraph(text)

    summary_rows: list[list[object]] = []
    short_rows: list[list[object]] = []
    long_rows: list[list[object]] = []
    metric_index: dict[tuple[str, str, str], list[float]] = {}

    for model in data.get("models", []):
        model_name = model.get("model", "")
        total = 0
        ok = 0
        worst = None
        for channel in model.get("channels", []):
            for scenario in channel.get("scenarios", []):
                scenario_name = scenario.get("scenario", {}).get("name", "")
                for level in scenario.get("levels", []):
                    total += int(level.get("total", 0))
                    ok += int(level.get("ok", 0))
                    row = [
                        model_name,
                        level.get("concurrency", ""),
                        f"{level.get('ok', 0)}/{level.get('total', 0)}",
                        f"{float(level.get('success_rate_pct', 0)):.1f}%",
                        fmt_ms(level.get("ttft", {}).get("avg_ms")),
                        fmt_ms(level.get("ttft", {}).get("p95_ms")),
                        fmt_ms(level.get("e2e", {}).get("avg_ms")),
                        fmt_ms(level.get("e2e", {}).get("p95_ms")),
                        fmt_1(level.get("tpot", {}).get("avg_ms")),
                        fmt_1(level.get("per_request_tok_per_s", {}).get("avg")),
                    ]
                    if "短文本" in scenario_name:
                        short_rows.append(row)
                    else:
                        long_rows.append(row)
                    kind = scenario_kind(scenario_name)
                    for metric_name, metric_value in [
                        ("首次生成token时间（TTFT）", level.get("ttft", {}).get("avg_ms", 0)),
                        ("每token延迟（TPOT）", level.get("tpot", {}).get("avg_ms", 0)),
                    ]:
                        metric_index.setdefault((kind, metric_name, model_name), []).append(float(metric_value or 0))
                    if worst is None or level.get("success_rate_pct", 0) < worst.get("success_rate_pct", 0):
                        worst = level
        success_rate = ok / total * 100 if total else 0
        summary_rows.append([
            model_name,
            f"{ok}/{total}",
            f"{success_rate:.1f}%",
            f"C={worst.get('concurrency')} / {float(worst.get('success_rate_pct', 0)):.1f}%" if worst else "-",
        ])

    metric_rows: list[list[object]] = []
    for kind, metric_name, desc in [
        ("短文本", "首次生成token时间（TTFT）", "首包延迟"),
        ("长文本", "首次生成token时间（TTFT）", "首包延迟"),
        ("短文本", "每token延迟（TPOT）", "输出平稳性"),
        ("长文本", "每token延迟（TPOT）", "输出平稳性"),
    ]:
        for model_name in models:
            vals = metric_index.get((kind, metric_name, model_name), [])
            metric_rows.append([
                f"{kind}-{metric_name}",
                "ms",
                desc,
                model_name,
                ", ".join(fmt_1(v) for v in vals if v) or "-",
                fmt_1(avg(vals)),
            ])

    doc.add_paragraph("性能测评关注指标：")
    add_table(doc, ["指标名称", "单位", "描述", "测评模型", "实测值", "平均值"], metric_rows)

    doc.add_heading("五、总体结论", level=1)
    add_table(doc, ["模型", "成功/总请求", "总成功率", "最低成功率档位"], summary_rows)
    doc.add_paragraph("低并发结果可用于判断模型基础可用性；高并发档位用于识别供应商限流、渠道容量和网关超时风险。")
    doc.add_paragraph("长文本仅覆盖单并发基线；因此不对长文本高并发稳定性作结论。")

    doc.add_heading("六、短文本并发测试结果", level=1)
    add_table(doc, ["模型", "并发", "OK/Total", "成功率", "TTFT avg(ms)", "TTFT p95(ms)", "Latency avg(ms)", "Latency p95(ms)", "TPOT avg(ms)", "tokens/s avg"], short_rows)

    doc.add_heading("七、长文本基线测试结果", level=1)
    add_table(doc, ["模型", "并发", "OK/Total", "成功率", "TTFT avg(ms)", "TTFT p95(ms)", "Latency avg(ms)", "Latency p95(ms)", "TPOT avg(ms)", "tokens/s avg"], long_rows)

    doc.add_heading("八、风险与建议", level=1)
    for text in [
        "若高并发成功率不足，建议排查 key 分组路由、供应商限流、渠道池容量和网关超时配置。",
        "若客户要求长文本高并发，需要提高长文本请求数；否则报告只能代表长文本单并发基线。",
        "建议对失败样本按 request_id 分析错误类型，区分供应商限流、网关超时、客户端超时和内容安全拒绝。",
    ]:
        doc.add_paragraph(text)

    return doc


def main() -> int:
    parser = argparse.ArgumentParser(description="Convert fy-poc-loadtest JSON to DOCX.")
    parser.add_argument("--input", required=True, help="Path to poc_loadtest_*.json")
    parser.add_argument("--output", required=True, help="Path to output .docx")
    args = parser.parse_args()

    data = json.loads(Path(args.input).read_text(encoding="utf-8"))
    doc = build_doc(data)
    out = Path(args.output)
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out)
    print(out)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
